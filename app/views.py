import os
import pandas as pd 
import numpy as np
from app import app
from io import BytesIO
import warnings
from warnings import *
from sqlalchemy import text
from flask import render_template,request, jsonify,flash,redirect,url_for,session,send_file
import xgboost as xgb
import pickle
import google.generativeai as genai
from dotenv import load_dotenv
from .models.db_config import *
import mysql.connector
import math
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
# database 
def create_db_connection():
    return mysql.connector.connect(
        host=DB_HOST,
        user=DB_USER,
        password=DB_PASSWORD,
        database=DB_NAME
    )

db_connection = create_db_connection()

# api gemini
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# load model xgboost
model = pickle.load(open('app/xgb_model.pkl', 'rb'))

# Function regis dan login 
def register_admin(username, password):
    try:
        conn = create_db_connection()
        cursor = conn.cursor()
        nextidAdmin = generate_next_id_admin()
        query = "INSERT INTO admin (id_admin,username, password) VALUES (%s, %s ,%s)"
        cursor.execute(query, (nextidAdmin,username, password))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print("Error in register_user:", e)
        return False

def generate_next_id_admin():
        cursor = db_connection.cursor()
        cursor.execute("SELECT MAX(id_admin) FROM admin") 
        max_id = cursor.fetchone()[0]

        if max_id:
            prefix = max_id[0] 
            number = int(max_id[1:]) 
            next_number = number + 1
            next_id = f"{prefix}{next_number:02d}"
        else:
            next_id = "A01" 

        return next_id
# Fungsi untuk melakukan login 
def login_admin(username, password):
    try:
        conn = create_db_connection()
        cursor = conn.cursor()
        query = "SELECT * FROM admin WHERE username = %s AND password = %s"
        cursor.execute(query, (username, password))
        user = cursor.fetchone()
        conn.close()
        return user
    except Exception as e:
        print("Error in login admin:", e)
        return None

# route
## Halaman registrasi
@app.route('/register', methods=['GET', 'POST'])
def register():
    success_message = None
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if register_admin(username, password):
            success_message = 'Registrasi sukses.'
        else:
            return jsonify({'error': 'Failed to register user.'}), 500
    return render_template('public/register.html', success_message=success_message)

# Halaman login
@app.route("/")
@app.route('/login', methods=['GET', 'POST'])
def login():
    warning_message = None
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = login_admin(username, password)
        if user:
            session['username'] = username
            # Redirect ke halaman beranda atau yang diinginkan setelah login berhasil
            return redirect(url_for('index'))
        else:
            warning_message = 'Username atau password salah. Silakan coba lagi atau daftar terlebih dahulu.'
            return render_template('public/login.html', warning_message=warning_message)
    return render_template('public/login.html', warning_message=warning_message)

# Fungsi untuk logout
@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))


@app.route("/home")
def index():
    if 'username' in session:
        return render_template('public/index.html')
    return redirect(url_for('login'))


@app.route("/about")
def about():
    return render_template('public/about.html')




@app.route("/riwayat")
def riwayat():
    page = request.args.get('page', 1, type=int)
    per_page = 5
    cursor = db_connection.cursor()

    # Query untuk menggabungkan tabel diagnosa dan riwayat
    query = f"""
        SELECT riwayat_diagnosa.id_riwayat, riwayat_diagnosa.id_diagnosa, diagnosa.Stage
        FROM riwayat_diagnosa
        INNER JOIN diagnosa ON diagnosa.id_diagnosa = riwayat_diagnosa.id_diagnosa
        ORDER BY riwayat_diagnosa.id_riwayat ASC
        LIMIT {per_page}
        OFFSET {(page-1)*per_page} 

    """
    cursor.execute(query)
    data_riwayat = cursor.fetchall()
    total_pages = get_total_pages(cursor)
    
    return render_template('public/riwayat.html',data_riwayat = data_riwayat, page=page, per_page=per_page, total_pages=total_pages)

def get_total_pages(cursor):
    per_page = 5
    # Query untuk menghitung total data
    cursor.execute("SELECT COUNT(*) FROM riwayat_diagnosa")
    total_data = cursor.fetchone()[0]

    # Hitung total halaman
    return int(math.ceil(total_data / per_page))

@app.route("/riwayat/delete/<id_riwayat>", methods=['POST'])
def delete_riwayat(id_riwayat):
    cursor = db_connection.cursor()
    cursor.execute("DELETE FROM riwayat_diagnosa WHERE id_riwayat = %s", (id_riwayat,))
    db_connection.commit()

    return redirect(url_for('riwayat'))

@app.route("/models")
def models():
    return render_template('public/models.html')

@app.route("/diagnosa")
def diagnosa():
    results = ''
    return render_template('public/diagnosa.html')


@app.route("/predict", methods=['GET', 'POST'])

def predict():
    avgBilirubin = 1.0
    avgCholes = 200.0
    avgAlbumin = (3.5,5.9)
    avgAlkPhos = (33,96)
    avgSGOT = (5,44)
    avgTryglicerides = 150
    avgPlatelets = (150,450)
    avgProthrombin = (11,13)

    Drug = float(request.form['Drug'])
    Age = float(request.form['Age'])
    Sex = float(request.form['Sex'])
    Ascites = float(request.form['Ascites'])
    Hepatomegaly = float(request.form['Hepatomegaly'])
    Spiders = float(request.form['Spiders'])
    Edema = float(request.form['Edema'])
    Bilirubin = float(request.form['Bilirubin'])
    Cholesterol = float(request.form['Cholesterol'])
    Albumin = float(request.form['Albumin'])
    Copper = float(request.form['Copper'])
    Alk_Phos = float(request.form['Alk_Phos'])
    SGOT = float(request.form['SGOT'])
    Tryglicerides = float(request.form['Tryglicerides'])
    Platelets = float(request.form['Platelets'])
    Prothrombin = float(request.form['Prothrombin'])
    Status = float(request.form['Status'])

    data = {
        'Drug': [Drug],
        'Age': [Age],
        'Sex': [Sex],
        'Ascites': [Ascites],
        'Hepatomegaly': [Hepatomegaly],
        'Spiders': [Spiders],
        'Edema': [Edema],
        'Bilirubin': [Bilirubin],
        'Cholesterol': [Cholesterol],
        'Albumin': [Albumin],
        'Copper': [Copper],
        'Alk_Phos': [Alk_Phos],
        'SGOT': [SGOT],
        'Tryglicerides': [Tryglicerides],
        'Platelets': [Platelets],
        'Prothrombin': [Prothrombin],
        'Status': [Status]
    }

    # Membuat DataFrame dari data
    df = pd.DataFrame(data)

    # Membuat objek DMatrix dari DataFrame
    dtest = xgb.DMatrix(df)

    # Memprediksi hasil
    results = model.predict(dtest)[0]

    def stageCondition():
        if results == 3.0:
            return 'Stage 4 [Liver Cancer]'
        elif results == 1.0:
            return 'Stage 2 [Fibrosis Liver]'
        elif results == 2.0:
            return 'Stage 3 [Cirrhosis Liver]'
        else:
            return 'Stage 1 [Fatty Cancer]'
    
    def solution(): 
        if results == 3.0:
            return 'Disarankan agar melakukan beberapa prosedur seperti <br><br> 1. Operasi merupakan salah satu pilihan pengobatan kanker hati. Prosedur ini melibatkan pengangkatan tumor atau bagian hati yang terkena kanker. Keputusan untuk melakukan operasi tergantung pada lokasi dan ukuran tumor, serta kondisi kesehatan pasien. Operasi dapat membantu mengurangi beban tumor dan memperbaiki fungsi hati.<br><br> 2. Ablasi adalah tindakan untuk menghancurkan kanker tanpa harus mengangkatnya. Metode ini digunakan pada kanker berukuran kecil, atau pada pasien yang tidak bisa menjalani operasi pengangkatan atau transplantasi hati.<br><br> 3. Embolisasi adalah penyuntikan obat untuk menghalangi atau mengurangi aliran darah ke sel kanker di hati. Prosedur ini dilakukan pada kanker yang berukuran lebih dari 5 cm dengan fungsi hati yang masih cukup baik.Kemoterapi <br><br> 4. Kemoterapi adalah pemberian obat-obatan untuk membunuh sel kanker. Prosedur ini digunakan pada kanker hati yang tidak dapat diangkat melalui operasi, dan pada kanker hati yang tidak bisa diatasi dengan ablasi, embolisasi, atau terapi target.<br><br> 5. Kemoembolisasi atau trans arterial chemoembolisation (TACE) adalah terapi yang memadukan embolisasi dengan kemoterapi. TACE dilakukan dengan menyuntikkan obat ke area kanker dan menghambat aliran darah ke sel kanker di hati. Kemoembolisasi dilakukan pada pasien yang tidak bisa menjalani prosedur operasi dan pada pasien yang sedang menunggu donor hati. Kemoembolisasi juga dapat dilakukan untuk menyusutkan tumor agar bisa diangkat melalui bedah.'
        elif results == 1.0:
            return 'Disarankan untuk berhenti minum alkohol dan juga mengurangi mengonsumsi makanan tinggi akan glukosa dan kolesterol. Lalu melakukan diet dan mengurangi berat badan, mengonsumsi makanan yang tinggi serat, protein, lemak sehat, kompleks karbohidrat, vitamin dan mineral. Suplemen omega-3 juga dapat bermanfaat untuk mengurangi peradangan dan meningkatkan fungsi hati.Selain itu penting untuk memperhatikan asupan cairan dengan cukup, terutama air, untuk membantu menjaga hidrasi tubuh dan fungsi hati yang optimal'
        elif results == 2.0:
            return 'Disarankan untuk mengonsumsi makanan rendah garam dan tablet spironolactone untuk mengurangi kelebihan cairan di dalam tubuh, serta mengonsumsi propranolol untuk mengurangi tekanan yang tinggi di dalam hati. Selain itu, suplemen juga dapat diberikan untuk mengatasi kekurangan nutrisi dan mencegah pengeroposan tulang, sementara krim dapat digunakan untuk mengatasi rasa gatal. Gastroskopi juga dapat dilakukan untuk mengikat pembuluh darah yang melebar di kerongkongan dan berisiko menimbulkan perdarahan. <br> Namun, selain meredakan gejala, penting untuk menangani penyebab yang mendasari sirosis. Hal ini meliputi menghentikan konsumsi minuman beralkohol, mengonsumsi obat antivirus untuk hepatitis, dan menurunkan berat badan pada pasien sirosis yang juga mengalami obesitas. Sebelum mengonsumsi obat-obatan apa pun, penderita sirosis disarankan untuk selalu berkonsultasi terlebih dahulu dengan dokter, karena beberapa obat dapat memperberat kerja organ hati dan memperburuk gejala sirosis. \n jika kondisi hati sudah tidak memungkinkan sebaiknya melakukan transplantasi hati karena Jaringan hati yang sudah menjadi jaringan parut tidak dapat kembali normal. Oleh karena itu, pasien yang kerusakan hatinya sudah parah dan fungsi hatinya sudah sangat menurun perlu menjalani cangkok hati atau transplantasi hati'
        else:
            return 'Untuk Alcoholic Liver Pasien disarankan melakukan Abstinensia alkohol, karena sangat penting untuk membatasi perkembangan steatosis dan mencegah kerusakan hati lebih lanjut. Diperlukan modifikasi gaya hidup seperti berhenti merokok, mengontrol diet, dan pengendalian berat badan.\n Disarankan juga untuk Pasien ALD direkomendasikan untuk mendapatkan asupan 1,2–1,5 gram protein/kgBB dan kalori 30–35 kkal/kgBB per hari, dengan peningkatan frekuensi makan termasuk makanan ringan pada malam hari <br> Suplementasi mikronutrien juga direkomendasikan.\n Untuk Non Alcoholic pasien adalah penurunan berat badan dan manajemen komorbiditas seperti obesitas, dislipidemia, resistensi insulin, dan diabetes tipe 2. Modifikasi gaya hidup termasuk perbaikan diet, olahraga, dan modifikasi perilaku dapat membantu memperlambat progresivitas Fatty liver. Penggunaan atorvastatin yang dikombinasikan dengan vitamin C dan E telah ditunjukkan dapat mengurangi risiko steatosis'
    def bilAvg():
        
        if Bilirubin >= avgBilirubin:
            return 'Tingkat Bilirubin anda Tinggi \n Konsumsi Buah dan Sayur, Hindari Alkohol dan Tetap terhidrasi'
        else:
            return 'Tingkat Bilirubin anda Normal'
    
    def CholAvg():
        if Cholesterol >= avgCholes:
            return 'Tingkat Kolesterol anda Tinggi \n Kurangi mengkonsumsi makanan berlemak, Hindari Alkohol dan Rutin Berolahraga'
        else:
            return 'Tingkat Cholesterol anda normal'
    
    def albAvg():
        if avgAlbumin[0] <= Albumin <= avgAlbumin[1]:
            return 'Tingkat Albumin anda Normal'
        else:
            if Albumin < avgAlbumin[0]:
                return 'Tingkat ALbumin anda rendah \n Konsumsi Makanan kaya akan protein seperti  tahu, tempe dan ikan '
            else:
                return 'Tingkat Albumin anda Tinggi \n Kurangi makanan berprotein dan konsumsi makanan tinggi serat'
    def alkAvg():
        if avgAlkPhos[0] <= Alk_Phos <= avgAlkPhos[1]:
            return 'Tingkat Alkaline Fosfatase anda Normal'
        else:
            if Alk_Phos < avgAlkPhos[0]:
                return 'Tingkat Alkaline Fosfatase anda rendah \n Konsumsi Makanan kaya akan kalsium dan fosfat seperti susu, kacang-kacangan, dan sayuran hijau \n Dan Konsultasi ke Dokter lebih lanjut'
            else:
                return 'Tingkat Alkaline Fosfatase anda tinggi \n Disarankan Konsultasi ke dokter dikarenakan bisa saja ada kerusakan pada hati, tulang atau kantong empedu \n Dan hindari mengonsumsi alkohol '
    
    def SGOTAvg():
        if avgSGOT[0] <= SGOT <= avgSGOT[1]:
            return 'Tingkat SGOT anda Normal'
        else:
            if SGOT < avgSGOT[0]:
                return 'Konsumsi Makanan kaya akan Vitamin B6 dan biji-bijian, \n dan hindari Alkohol dan juga obat-obatan beracun yang dapat merusak hati'
            else:
                return 'Konsultasikan ke dokter karena penyebab SGOT tinggi bisa jadi macam macam \n dan hindari alkohol dan obat obatan beracun bagi organ dalam'
    def TryAvg():
        if Tryglicerides >= avgTryglicerides:
            return 'Tryglicerides anda tinggi  \n Kurangi konsumsi makanan karbohidrat olahan seperti pasta, mie. dan makan makanan penurun tryglicerides  seperti salmon, ikan cod, dan tuna. Ikan berlemak yang mengandung asam lemak omega 3 '
        else:
            return 'Tryglicerides anda rendah \n Trigliserida rendah dapat disebabkan oleh beberapa faktor, seperti pola makan rendah lemak, kekurangan gizi, atau efek samping obat. Untuk mengatasinya, Anda perlu mengonsumsi makanan yang lebih banyak mengandung lemak sehat, seperti alpukat, kacang-kacangan, dan minyak zaitun'
    def PlatAvg():
        if avgPlatelets[0] <= Platelets <= avgPlatelets[1]:
            return 'Tingkat Platelets/trombosit anda Normal'
        else:
            if Platelets < avgPlatelets[0]:
                return 'Tingkat Platelets/trombosit anda rendah \n Konsumsi sumber vitamin C, seperti jeruk, kiwi, brokoli, dan paprika. Vitamin C merupakan jenis vitamin yang dapat membantu proses penyerapan zat besi sehingga bisa mengoptimalkan proses pembentukan sel darah, termasuk trombosit.'
            else:
                return 'Tingkat Platelets/trombosit anda tinggi \n Konsumsi makanan penurun trombosit, seperti bawang putih mentah, buah delima, dan makanan kaya vitamin E. Hindari makanan berlemak dan berminyak, seperti gorengan dan makanan cepat saji'
    def ProthAvg():
        if avgProthrombin[0] <= Prothrombin <= avgProthrombin[1]:
            return 'Tingkat Platelets/trombosit anda Normal'
        else:
            if Prothrombin < avgProthrombin[0]:
                return 'Tingkat Platelets/trombosit anda rendah \n Konsumsi sumber vitamin C, seperti jeruk, kiwi, brokoli, dan paprika. Vitamin C merupakan jenis vitamin yang dapat membantu proses penyerapan zat besi sehingga bisa mengoptimalkan proses pembentukan sel darah, termasuk trombosit.'
            else:
                return 'Tingkat Platelets/trombosit anda tinggi \n Konsumsi makanan penurun trombosit, seperti bawang putih mentah, buah delima, dan makanan kaya vitamin E. Hindari makanan berlemak dan berminyak, seperti gorengan dan makanan cepat saji'
    
    def allSolution():
        stage = stageCondition()
        solusi = solution()
        bilirubinAvg = bilAvg()
        CholesAvg =CholAvg()
        albuminAvg = albAvg()
        alkPhosAvg = alkAvg()
        SGOTAverage = SGOTAvg()
        TrygliAvg = TryAvg()
        plateletsAvg = PlatAvg()
        prothrombinAvg =ProthAvg()

        # Menggabungkan semua hasil diagnosa menjadi satu teks
        all_solution = f"{stage}, {solusi}, Bilirubin: {bilirubinAvg}, Cholesterol: {CholesAvg},Albumin: {albuminAvg}, Alkaline Phosphatase: {alkPhosAvg}, SGOT: {SGOTAverage}, Triglyceride: {TrygliAvg}, Platelets: {plateletsAvg}, Prothrombin: {prothrombinAvg}"
        return all_solution

    # Fungsi untuk membuat ID berikutnya berdasarkan ID terakhir yang ada di tabel
    def generate_next_id_diagnosa():
        cursor = db_connection.cursor()
        cursor.execute("SELECT MAX(id_diagnosa) FROM diagnosa") 
        max_id = cursor.fetchone()[0]

        if max_id:
            prefix = max_id[0] 
            number = int(max_id[1:]) 
            next_number = number + 1
            next_id = f"{prefix}{next_number:02d}"
        else:
            next_id = "D01" 

        return next_id
    
    def generate_next_id_riwayat():
        cursor = db_connection.cursor()
        cursor.execute("SELECT MAX(id_riwayat) FROM riwayat_diagnosa") 
        max_id = cursor.fetchone()[0]

        if max_id:
            prefix = max_id[0] 
            number = int(max_id[1:]) 
            next_number = number + 1
            next_id = f"{prefix}{next_number:02d}"
        else:
            next_id = "R01" 

        return next_id

    def inHistory():
        if request.method == 'POST':
            next_idDiagnosa = generate_next_id_diagnosa()
            nextidRiwayat = generate_next_id_riwayat()
            dataDiagnosa = np.array([ Drug, Age, Sex, Ascites, Hepatomegaly, Spiders, Bilirubin, Cholesterol, Albumin, Copper, Alk_Phos, SGOT, Tryglicerides, Platelets, Prothrombin, Edema, Status, results],dtype=np.float64)
            dataRiwayat = allSolution()
            dataDiagnosa = [float(x) for x in dataDiagnosa]
            data_list = list(dataDiagnosa)


            # Simpan data ke database
            cursor = db_connection.cursor()
            
            cursor.execute("INSERT INTO diagnosa (id_diagnosa, Drugs, Age, Sex, Ascites, Hepatomegaly, Spiders, Bilirubin, Cholesterol, Albumin, Copper, Alk_Phos, SGOT, Tryglicerides, Platelets, Prothrombin, Edema, Status, Stage) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                        (next_idDiagnosa,*data_list))
            cursor.execute("INSERT INTO riwayat_diagnosa (id_riwayat,id_diagnosa,solusi_diagnosa) VALUES (%s, %s, %s)", (nextidRiwayat,next_idDiagnosa,dataRiwayat))
            db_connection.commit()
    
    return render_template('public/predict.html', results=results, stageCondition = stageCondition(), bilAvg =bilAvg(), Cholavg=CholAvg(), albAvg =albAvg(), alkAvg=alkAvg(), SGOTAvg = SGOTAvg(), TryAvg=TryAvg(), PlatAvg=PlatAvg(), ProthAvg=ProthAvg(), solution =solution(),allSolution=allSolution(),inHistory=inHistory())

@app.route("/print_diagnosa", methods=['GET'])
def print_diagnosa():
    # Ambil data dari dua tabel dan lakukan join
    id_riwayat = request.args.get('id_riwayat')
    cursor = db_connection.cursor()
    cursor.execute("SELECT d.*, r.solusi_diagnosa FROM diagnosa d JOIN riwayat_diagnosa r ON d.id_diagnosa = r.id_diagnosa WHERE r.id_riwayat = %s", (id_riwayat,))
    data_diagnosa_print = cursor.fetchone()

    # Check if data is fetched correctly
    if not data_diagnosa_print:
        return "No data found", 404

    # Ensure each row has expected structure
    if len(data_diagnosa_print) != 20:
        return "Unexpected data structure", 500

    # Buat dokumen Word baru
    document = Document()
    document.add_heading('Hasil Diagnosa Tahapan Penyakit ', 0)
    document.add_paragraph()
    document.add_heading('Keterangan Attribut',level=2)
    document.add_paragraph('Drugs 1.0 = Iya, dan 0 = Tidak')
    document.add_paragraph('Gender 1.0 = Pria, dan 0 = Wanita')
    document.add_paragraph('Ascites 1.0 = Iya, dan 0 = Tidak')
    document.add_paragraph('Hepatomegaly 1.0 = Iya, dan 0 = Tidak')
    document.add_paragraph('Spiders 1.0 = Iya, dan 0 = Tidak')
    document.add_paragraph('Edema 2.0 = Iya dengan Diuretik, 1.0 = Iya Tanpa Diuretik, dan 0 = Tidak')
    document.add_paragraph('Status 2.0 = Death,1.0 = CL(Censored due to Liver), dan 0 = C(Censored)')
    document.add_paragraph('Status 3.0 = Liver Cancer, 2.0 = Cirrhosis Liver,1.0 = Fibrosis Liver, dan 0 = Fatty Liver')
    # Tambahkan baris kosong setelah heading
    document.add_paragraph()

    # Add data to the document
    document.add_heading(f"ID Diagnosa: {data_diagnosa_print[0]}", level=3)
    document.add_paragraph()
    document.add_paragraph(f"Drugs: {data_diagnosa_print[1]}", style='List Bullet')
    document.add_paragraph(f"Umur: {data_diagnosa_print[2]}", style='List Bullet')
    document.add_paragraph(f"Gender: {data_diagnosa_print[3]}", style='List Bullet')
    document.add_paragraph(f"Ascites: {data_diagnosa_print[4]}", style='List Bullet')
    document.add_paragraph(f"Hepatomegaly: {data_diagnosa_print[5]}", style='List Bullet')
    document.add_paragraph(f"Spiders: {data_diagnosa_print[6]}", style='List Bullet')
    document.add_paragraph(f"Bilirubin: {data_diagnosa_print[7]}", style='List Bullet')
    document.add_paragraph(f"Cholesterol: {data_diagnosa_print[8]}", style='List Bullet')
    document.add_paragraph(f"Albumin: {data_diagnosa_print[9]}", style='List Bullet')
    document.add_paragraph(f"Copper: {data_diagnosa_print[10]}", style='List Bullet')
    document.add_paragraph(f"Alk_Phos: {data_diagnosa_print[11]}", style='List Bullet')
    document.add_paragraph(f"SGOT: {data_diagnosa_print[12]}", style='List Bullet')
    document.add_paragraph(f"Tryglicerides: {data_diagnosa_print[13]}", style='List Bullet')
    document.add_paragraph(f"Trombosit: {data_diagnosa_print[14]}", style='List Bullet')
    document.add_paragraph(f"Prothrombin: {data_diagnosa_print[15]}", style='List Bullet')
    document.add_paragraph(f"Edema: {data_diagnosa_print[16]}", style='List Bullet')
    document.add_paragraph(f"Status: {data_diagnosa_print[17]}", style='List Bullet')
    document.add_paragraph(f"Stage: {data_diagnosa_print[18]}", style='List Bullet')
    document.add_paragraph()
    document.add_paragraph(f"Solusi: {data_diagnosa_print[19]}")

    # Simpan dokumen Word ke dalam BytesIO
    output = BytesIO()
    document.save(output)
    output.seek(0)

    # Kirim file Word sebagai respons
    return send_file(output, as_attachment=True, download_name="hasil_diagnosa.docx")

